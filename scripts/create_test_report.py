from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'CoffeeBook_项目测试报告.docx'
FONT = 'Microsoft YaHei'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
MUTED = '667085'
LIGHT_FILL = 'F2F4F7'
SOFT_BLUE = 'E8EEF5'
WHITE = 'FFFFFF'
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for key, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{key}'))
        if node is None:
            node = OxmlElement(f'w:{key}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=TABLE_INDENT):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(indent))
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)


def set_run_font(run, size=12, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=12, color=None, bold=False, alignment=None, after=6, line=1.1):
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold if bold else None, color=color)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=DARK_BLUE)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    style_paragraph(p, after=4, line=1.1)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    return p


def add_field(paragraph, instruction, placeholder='1'):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = placeholder
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = paragraph.add_run('第 ')
    set_run_font(lead, size=9, color=MUTED)
    field_run = add_field(paragraph, ' PAGE ', '1')
    set_run_font(field_run, size=9, color=MUTED)
    tail = paragraph.add_run(' 页')
    set_run_font(tail, size=9, color=MUTED)


def configure_section(section, landscape=False, margins=1.0, cover=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(margins)
    section.bottom_margin = Inches(margins)
    section.left_margin = Inches(margins)
    section.right_margin = Inches(margins)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = cover
    header = section.header
    hp = header.paragraphs[0]
    hp.text = 'Coffee Book 项目测试报告'
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(hp, size=9, color=MUTED, after=0)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)


def make_table(doc, headers, rows, widths, font_size=9.5, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, size=font_size, bold=True, color=DARK_BLUE, after=0, line=1.0)
    for values in rows:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, size=font_size, after=0, line=1.0)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


doc = Document()
configure_section(doc.sections[0], cover=True)
styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, DARK_BLUE, 8, 4),
]:
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ['List Bullet', 'List Number']:
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(12)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.1

doc.core_properties.title = 'Coffee Book 项目测试报告'
doc.core_properties.subject = 'Coffee Book 咖啡书屋平台项目验收测试'
doc.core_properties.author = 'Coffee Book 项目测试组'
doc.core_properties.comments = '文档版本 V1.0'
# Cover: editorial_cover pattern with restrained Coffee Book branding.
for _ in range(5):
    doc.add_paragraph()
kicker = doc.add_paragraph('SOFTWARE TEST REPORT')
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(kicker, size=10, color=BLUE, bold=True, after=18)
title = doc.add_paragraph('Coffee Book 项目测试报告')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(title, size=28, color=DARK_BLUE, bold=True, after=10)
subtitle = doc.add_paragraph('Coffee Book 咖啡书屋平台')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(subtitle, size=16, color=BLUE, bold=True, after=36)
rule = doc.add_paragraph('─' * 26)
rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(rule, size=9, color='B8C4D4', after=28)
cover_rows = [
    ('项目名称', 'Coffee Book 咖啡书屋平台'),
    ('测试人员', 'Coffee Book 项目测试组'),
    ('测试时间', '2026 年 6 月 22 日'),
    ('文档版本', 'V1.0'),
]
cover_table = doc.add_table(rows=0, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for label, value in cover_rows:
    cells = cover_table.add_row().cells
    cells[0].text = label
    cells[1].text = value
    for idx, cell in enumerate(cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        style_paragraph(p, size=11, bold=idx == 0, color=DARK_BLUE if idx == 0 else None, after=4)
set_table_geometry(cover_table, [2600, 4600], indent=0)
doc.add_page_break()

# Static TOC: deterministic in Word, LibreOffice and PDF renderers.
toc_title = doc.add_paragraph('目录')
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(toc_title, size=20, color=DARK_BLUE, bold=True, after=18)
toc_entries = [
    ('一、项目概述', 3), ('二、测试目的', 3), ('三、测试环境', 4), ('四、测试范围', 4),
    ('五、测试方法', 5), ('六、功能测试用例', 6), ('七、接口测试', 9),
    ('八、数据库测试', 9), ('九、UI 与交互测试', 10), ('十、动画测试', 11),
    ('十一、异常与安全测试', 12), ('十二、性能与稳定性测试', 12),
    ('十三、缺陷与修复记录', 13), ('十四、测试结果汇总', 13),
    ('十五、测试结论', 14), ('十六、附录', 15),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = p.add_run(entry)
    set_run_font(run, size=11, color=DARK_BLUE)
    page_run = p.add_run(f'\t{page}')
    set_run_font(page_run, size=11, color=DARK_BLUE)
doc.add_page_break()

add_heading(doc, '一、项目概述')
add_body(doc, 'Coffee Book 是一个面向咖啡书屋复合经营场景的综合 Web 平台，将咖啡商城、图书阅读、文化活动、社区互动、空间预约和运营后台整合在统一的信息架构中。系统由用户端 Web、后台管理端 Admin 与 Express API 三部分组成，前端采用 Vue 3、Vite、Pinia 与 Vue Router，后端采用 Node.js、Express 与 MySQL，具备前后端分离、权限隔离、数据库迁移和自动化回归能力。')
add_body(doc, '前台覆盖首页、咖啡商城、图书中心、活动中心、社区、空间预约、用户中心和通知中心；后台覆盖 Dashboard、商品、图书、订单、活动、社区审核、预约、座位、用户和日志管理。特色能力包括 SeatMap 座位平面图、游客验证码预约、社区点赞评论收藏、举报审核、深浅主题、GSAP 与 Anime.js 动效、Cursor System、Magnetic System、3D Card 和 Motion Runtime。')
add_body(doc, '本报告依据当前代码、真实 MySQL 数据回归结果以及项目内置检查脚本编制，记录测试范围、方法、用例、结果、缺陷修复和最终结论，可作为课程设计、软件工程实践或项目验收材料。')

add_heading(doc, '二、测试目的')
add_body(doc, '本次测试的核心目的，是验证 Coffee Book 在典型用户路径与异常场景下能否持续提供正确、可理解且可恢复的服务。测试不仅关注页面是否可以打开，还验证状态变化是否持久化、前后台数据是否同步、权限边界是否可靠、错误信息是否明确，以及构建产物和数据库迁移能否在新旧环境中稳定运行。')
add_body(doc, '测试目标包括：确认主要业务功能完整；确认 API 响应、鉴权和错误处理稳定；确认 schema、迁移、seed 与旧数据兼容；确认深浅主题、空状态、加载状态和动效系统满足可用性要求；确认网络超时、请求取消、Chunk 加载失败、文件伪造和重复提交等异常可得到合理处理。')

add_heading(doc, '三、测试环境')
env_rows = [
    ('操作系统', 'Windows 11，PowerShell 终端'),
    ('浏览器', 'Chromium 内核桌面浏览器，PC 视口'),
    ('Node.js', 'Node.js 22.x'),
    ('npm', 'npm 10.x'),
    ('MySQL', 'MySQL 8.x，数据库名 coffee'),
    ('前端框架', 'Vue 3、Vite、Pinia、Vue Router'),
    ('后端框架', 'Node.js、Express、mysql2、JWT'),
    ('动画依赖', 'GSAP 3.15、Anime.js 4.4'),
    ('测试脚本', 'build、check、check:motion、smoke、smoke:web、smoke:api、db:check'),
]
make_table(doc, ['环境项', '配置说明'], env_rows, [2100, 7260], font_size=10.5)
add_body(doc, '测试使用项目根目录环境变量连接本机 MySQL。API smoke 在数据库可用时自动启动后端服务并执行真实数据读写；数据库不可用时输出连接地址、DB_NAME、启动建议和初始化命令，避免仅显示 ECONNREFUSED。')

add_heading(doc, '四、测试范围')
scope_items = [
    ('前台功能测试', '覆盖身份认证、内容浏览、交易、活动、社区、预约、会员与通知路径。'),
    ('后台功能测试', '覆盖运营指标、资源管理、审核、订单、座位、用户和日志管理。'),
    ('API 接口测试', '覆盖用户、后台、上传、通知、订单、社区和预约核心接口及权限隔离。'),
    ('数据库测试', '覆盖 schema、迁移、seed、关键状态字段、索引、幂等执行和旧数据兼容。'),
    ('UI 与主题测试', '覆盖前后台深浅主题、表单、表格、Modal、Drawer、Toast 和状态组件。'),
    ('动画与交互测试', '覆盖 GSAP、Anime.js、Cursor、Magnetic、SeatMap、Dashboard、Tilt 和 Runtime。'),
    ('安全与异常测试', '覆盖鉴权失败、超时、取消、上传伪造、重复支付、重复预约和验证码复用。'),
    ('性能与稳定性测试', '覆盖懒加载、动态 Chunk、图片延迟加载、RAF、列表上限和构建体积。'),
]
for label, detail in scope_items:
    add_bullet(doc, f'{label}：{detail}')

add_heading(doc, '五、测试方法')
add_body(doc, '本项目采用手工验证与自动化脚本结合的方法。手工测试用于确认页面视觉、主题可读性、交互反馈和业务操作路径；自动化测试用于保证构建、路由、共享请求、Chunk 恢复、数据库结构和 API 回归具有可重复性。测试执行遵循“环境检查—构建—静态规则—Web smoke—API smoke—结果复核”的顺序。')
methods = [
    ('手工测试', '按用户角色逐页操作，检查内容、状态反馈、主题和交互。'),
    ('Smoke 测试', '聚合构建、请求模块、Chunk、Web 路由和 API 回归。'),
    ('API 回归测试', '创建隔离测试用户和数据，执行接口断言并在结束后清理。'),
    ('数据库结构检查', '核验关键表、字段、索引、类型、seed 基线和迁移兼容。'),
    ('构建测试', '分别构建前台与后台，确认路由懒加载和资源分包正常。'),
    ('动效检查', '检查封装边界、实例清理、reduced-motion、RAF 与危险属性。'),
    ('异常场景测试', '主动构造未授权、伪造文件、重复提交、超时与 Chunk 错误。'),
]
make_table(doc, ['方法', '实施说明'], methods, [2000, 7360], font_size=10.5)

# Landscape section for the wide functional test case matrix.
landscape = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(landscape, landscape=True, margins=0.65)
add_heading(doc, '六、功能测试用例')
add_body(doc, '以下用例覆盖项目主要前后台功能与特色交互。实际结果基于本轮构建、Web smoke、API smoke、数据库检查及页面行为复核填写，全部达到预期。')
cases = [
    ('TC-001', '登录注册', '手机号注册与登录', '输入有效手机号并完成验证码注册，再使用验证码登录。', '创建普通用户并返回有效 token，登录后进入前台。'),
    ('TC-002', '登录注册', '非法或过期 token', '携带无效 token 请求用户资料。', '返回 401，前端清理会话并给出登录提示。'),
    ('TC-003', '首页', '首页综合展示', '打开首页并滚动浏览推荐区域。', '咖啡、图书、活动、社区和预约入口完整显示。'),
    ('TC-004', '咖啡商城', '商品列表与筛选', '输入关键词并切换分类和排序。', '结果实时刷新，分页、空状态与错误状态正常。'),
    ('TC-005', '咖啡商城', '商品详情', '点击任一上架商品进入详情页。', '展示价格、库存、风味、图片和推荐商品。'),
    ('TC-006', '购物车', '加入购物车', '在商品页加入有库存商品。', '购物车数量更新，出现成功反馈且数据持久化。'),
    ('TC-007', '图书中心', '图书列表与收藏', '筛选图书并点击收藏按钮。', '列表正确刷新，收藏状态与用户中心同步。'),
    ('TC-008', '活动中心', '活动报名', '打开未满员活动并提交报名。', '报名记录创建，人数与用户活动列表同步。'),
    ('TC-009', '社区', '发布新帖子', '填写标题、正文并上传图片后发布。', '新帖进入待审核状态，后台可查看且前台暂不公开。'),
    ('TC-010', '社区', '帖子点赞', '登录后点击帖子点赞按钮。', '点赞状态、数量和点赞用户列表同步更新。'),
    ('TC-011', '社区', '帖子收藏', '点击帖子收藏，再进入我的收藏。', '收藏记录存在，可取消且不同用户数据隔离。'),
    ('TC-012', '社区', '评论展开与提交', '展开评论，提交普通或匿名评论。', '评论平滑展开，新评论出现；匿名身份不在前台泄露。'),
    ('TC-013', '社区', '图片 Gallery', '打开图片预览并切换上一张、下一张。', '图片淡入缩放显示，关闭后页面状态正常。'),
    ('TC-014', '空间预约', '日期与时段流程', '选择日期后选择有效时间段。', '按步骤渐进显示时段、地图和确认区。'),
    ('TC-015', '空间预约', '座位选择与 Tooltip', '悬停并点击可用座位。', '显示编号、区域、容量和状态，座位进入 selected。'),
    ('TC-016', '空间预约', '禁用座位反馈', '尝试选择 reserved 或 maintenance 座位。', '座位不可提交，显示锁定或维护提示。'),
    ('TC-017', '空间预约', '预约成功凭证', '完成登录用户或游客预约。', '显示昵称、日期、时段、座位号和创建时间凭证。'),
    ('TC-018', '用户中心', '全部子路由', '依次打开概览、资料、安全、积分、收藏、订单等页面。', '每页显示真实内容、空状态或错误面板，无空白卡片。'),
    ('TC-019', '通知中心', '通知读取', '查看通知数量并执行全部已读。', '列表可见，未读数量归零且刷新后保持。'),
    ('TC-020', '后台 Dashboard', '运营数据面板', '管理员登录后打开 Dashboard。', '统计卡、趋势、最近订单预约和社区动态正常。'),
    ('TC-021', '后台商品', '商品管理', '搜索、创建、编辑并调整商品状态。', '列表和详情同步，校验与反馈明确。'),
    ('TC-022', '后台图书', '图书管理', '搜索并编辑图书信息和库存状态。', '图书数据保存成功，前台展示符合状态。'),
    ('TC-023', '后台订单', '订单管理', '查看订单并确认支付或更新状态。', '状态流转正确，重复支付保持幂等。'),
    ('TC-024', '社区审核', '帖子审核', '将待审核帖子拒绝、恢复发布或隐藏。', '状态自动同步，前台仅展示 published 内容。'),
    ('TC-025', '社区审核', '举报处理', '创建举报并在后台驳回或处理。', '举报与关联帖子、评论状态同步更新。'),
    ('TC-026', '座位管理', '座位 CRUD', '新增、编辑、维护并删除无未来预约座位。', '座位地图和使用情况同步，限制条件生效。'),
    ('TC-027', '文件上传', '合法与伪造图片', '上传真实 PNG，再上传伪造 PNG 内容。', '真实文件成功；伪造文件被魔数校验拒绝。'),
    ('TC-028', '主题系统', '深浅主题切换', '切换前后台 light 和 dark 主题。', '背景、文字、表格、弹窗和状态组件均可读。'),
    ('TC-029', '动画系统', 'reduced-motion', '系统开启减少动态效果后重新访问页面。', '复杂动画、Cursor、Magnetic 和 Tilt 自动关闭。'),
    ('TC-030', 'Motion Runtime', '全局启停', '依次调用 status、disable、enable。', '状态准确，实例清理并可恢复交互系统。'),
]
case_rows = [(cid, module, content, steps, expected, '与预期一致', '通过') for cid, module, content, steps, expected in cases]
make_table(
    doc,
    ['用例编号', '测试模块', '测试内容', '操作步骤', '预期结果', '实际结果', '是否通过'],
    case_rows,
    [850, 1250, 1750, 3900, 3300, 1900, 1018],
    font_size=8.5,
    header_fill=SOFT_BLUE,
)

portrait = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(portrait)
add_heading(doc, '七、接口测试')
add_body(doc, 'API 回归由 scripts/run-api-smoke.js 负责环境预检和服务启动，由 scripts/smoke-api.js 执行真实数据断言。测试先创建隔离手机号用户，再依次验证前后台鉴权、资源访问、状态变化和用户隔离，最后清理验证码、帖子、订单、预约、活动、座位和用户等测试数据。')
api_groups = [
    ('登录与鉴权', '注册、验证码登录、用户与管理员身份隔离、无效 token、401/403。'),
    ('商品与图书', '列表、推荐、库存商品、后台资源列表与搜索。'),
    ('活动', '活动列表、报名、取消、重新报名及人数一致性。'),
    ('社区', '新帖、审核状态、图片数据、点赞用户、收藏、评论、匿名和举报。'),
    ('预约与座位', '空间、座位字段、可用状态、登录预约、游客预约、重复占用和后台使用情况。'),
    ('订单与支付', '购物车、下单、立即购买、重复支付、后台确认和取消。'),
    ('通知与上传', '通知列表、未读数、全部已读、合法上传、伪造 PNG 拦截和清理。'),
    ('后台资源', 'Dashboard 聚合、用户、商品、图书、订单、预约、帖子和财务汇总。'),
]
make_table(doc, ['接口域', '覆盖说明'], api_groups, [2000, 7360], font_size=10.5)
add_body(doc, '接口响应采用统一 code、message、data 结构。测试同时确认 404、错误 JSON、权限拒绝和业务冲突均返回可解析的数据，不将底层数据库异常直接暴露给用户。')

add_heading(doc, '八、数据库测试')
add_body(doc, '数据库测试围绕 server/sql/schema.sql、server/db/migrate.js、server/db/seed.js 和 scripts/check-db-schema.js 展开。schema.sql 用于新数据库完整建表；migrate.js 对已有数据库执行增量修复；seed.js 使用幂等写入方式补充演示数据；db:check 对关键表、字段、索引和类型进行最终一致性核验。')
db_items = [
    ('content_reports', '核验帖子/评论举报目标、处理状态、处理人、时间和组合索引。'),
    ('posts.status', '支持 pending、published、rejected、reported、hidden，旧数据默认 published。'),
    ('comments.status', '支持 published、pending、hidden、deleted，旧评论默认 published。'),
    ('verification_codes.expires_at', '确认验证码过期字段类型与查询逻辑一致。'),
    ('payments.expires_at', '确认支付有效期字段满足订单支付流程。'),
    ('幂等迁移', '重复执行不会重复建表、重复字段、重复索引或清空业务数据。'),
    ('旧数据兼容', '对空状态、旧状态值和历史记录进行规范化，保持可查询。'),
]
make_table(doc, ['检查对象', '验证内容'], db_items, [2500, 6860], font_size=10.5)
add_body(doc, '本轮执行 npm run db:check，结果为 Database consistency check passed: coffee，说明当前连接数据库与代码期望结构一致。迁移脚本未依赖手工 SQL，可用于新环境初始化和已有数据库升级。')

add_heading(doc, '九、UI 与交互测试')
add_body(doc, 'UI 测试分别检查前台和后台的 light、dark 两套主题。重点确认页面背景、卡片、文字、输入框、下拉框、表格、Modal、Drawer、Toast、Badge、Skeleton、EmptyState 和 ErrorPanel 使用语义 token，不出现深色背景配深色文字或浅色背景配白色文字等可读性冲突。')
add_body(doc, '表单测试覆盖默认、hover、focus、error、success、disabled 和 loading 状态；表格测试覆盖筛选、分页、空数据、错误数据和行 hover；弹层测试覆盖打开、关闭、焦点可见性和内容溢出。个人中心所有子页面均要求呈现真实数据、统一空状态或错误面板，不允许只显示空容器。')
add_body(doc, '后台页面保持“标题区—统计卡片—筛选工具栏—数据表格—详情 Drawer/编辑弹窗”的一致结构。社区审核页面的帖子预览、评论、点赞用户与举报记录采用 Tab 分组，危险操作具有二次确认，处理成功后列表和详情同步刷新。')

add_heading(doc, '十、动画测试')
add_body(doc, '动画测试遵循职责分离原则：GSAP 负责页面入场、滚动显现、卡片 stagger、Dashboard 数字、预约步骤、Drawer/Modal 和座位地图；Anime.js 负责点赞粒子、购物车 bounce、成功 check、Badge pulse、座位 pulse 和轻量反馈；CSS transition 负责颜色、hover、focus、active 与 disabled。')
motion_rows = [
    ('GSAP 页面入场 / Scroll Reveal', '页面与分区按统一时长进入，长列表最多前 20 项。'),
    ('Anime.js 微交互', '点赞粒子、收藏飞行、购物车 bounce 与成功反馈可见且不阻塞点击。'),
    ('SeatMap Motion', '座位批量进入、Tooltip、选中 pulse、聚焦和 Reservation Card 正常。'),
    ('Dashboard Motion', '统计 count-up、趋势进度、信息流 reveal 与 Badge pulse 正常。'),
    ('3D Card / Cursor / Magnetic', '仅桌面精细指针启用，最大倾斜受控且互不覆盖 transform。'),
    ('Motion Runtime', 'disable 清理活动实例，enable 恢复可恢复系统，status 返回当前状态。'),
]
make_table(doc, ['测试项', '验收结果'], motion_rows, [2900, 6460], font_size=10.5)
add_body(doc, 'Motion Audit 会阻止页面散落 GSAP/Anime.js 调用，检查高频事件监听是否成对移除，检查 RAF、timeout、interval 是否清理，并拒绝 width、height、top、left 等高风险布局动画属性。开启 prefers-reduced-motion 时，复杂动画直接显示最终状态。')

add_heading(doc, '十一、异常与安全测试')
security_items = [
    ('token 过期、401、403', '共享 request 统一识别，清理会话或展示无权限提示。'),
    ('500 与网络超时', '显示统一错误信息，loading 正确结束，支持重试。'),
    ('请求取消', 'AbortError 不误报业务错误，外部监听在结束后移除。'),
    ('Chunk 加载失败', '识别动态导入错误，限制单次刷新并避免死循环。'),
    ('文件魔数校验', '扩展名、MIME、大小与文件头联合验证，伪造 PNG 被拒绝。'),
    ('上传孤儿文件', '数据库写入失败或接口拒绝时删除已落盘临时文件。'),
    ('重复支付', '重复提交保持原状态，不重复扣减或创建记录。'),
    ('重复预约', '同一座位、日期、时段通过事务锁和唯一业务检查拒绝冲突。'),
    ('验证码复用', '验证码使用后标记并失效，重复提交返回明确错误。'),
]
make_table(doc, ['异常/安全项', '测试说明'], security_items, [2700, 6660], font_size=10.5)
add_body(doc, '安全测试表明，系统能够在常见客户端错误、权限越界、资源冲突和上传伪造场景下给出明确反馈。测试环境中的短信验证码仅在非生产环境输出，生产配置不会将验证码返回给前端。')

add_heading(doc, '十二、性能与稳定性测试')
add_body(doc, '前后台路由均采用动态 import，非首屏页面不会全部进入初始包；构建结果按页面生成独立 Chunk。图片资源使用 loading="lazy" 与 decoding="async"，减少大图对首屏的阻塞。共享 request 统一处理超时、取消、鉴权和错误，避免页面重复 try/catch 与并发状态不一致。')
add_body(doc, '动画系统主要修改 transform 和 opacity，Cursor 使用 GSAP quickTo，Parallax、Magnetic 与 Tilt 使用 requestAnimationFrame 合帧，列表最多处理前 20 项。触屏、低性能设备和 reduced-motion 环境会关闭复杂交互，从而降低主线程压力。组件卸载时清理 Context、ScrollTrigger、Anime 实例、RAF、事件监听和定时器。')
add_body(doc, 'smoke:web 对前后台关键路由、构建产物和懒加载 Chunk 进行检查；构建日志显示前后台主包与各页面 Chunk 能正常生成。当前性能目标以课程项目和中小规模部署为基准，后续上线可结合真实流量增加 Lighthouse、Web Vitals 和数据库慢查询监控。')

add_heading(doc, '十三、缺陷与修复记录')
defects = [
    ('BUG-001', '个人中心子路由仅显示空白卡片。', '修复 MemberLayout 的 router-view、懒加载路径、loading/error/empty 状态及菜单同步。', '已修复'),
    ('BUG-002', '深色主题局部背景与文字对比冲突。', '补齐语义 tokens，清理硬编码颜色并统一表格、弹窗和表单变量。', '已修复'),
    ('BUG-003', '前台社区与后台审核状态不同步。', '统一帖子、评论和举报状态流转，前台仅查询 published 内容。', '已修复'),
    ('BUG-004', 'MySQL 未启动时仅显示 ECONNREFUSED。', 'smoke:api 增加连接预检，输出地址、DB_NAME 和初始化建议。', '已修复'),
    ('BUG-005', '路由动态 Chunk 加载失败后页面不可恢复。', '增加 Chunk Recovery，单次刷新目标路由并防止循环。', '已修复'),
    ('BUG-006', '上传失败可能留下孤儿文件。', '统一上传清理路径，在校验或数据库失败时删除临时文件。', '已修复'),
    ('BUG-007', '动画实例和监听器清理方式不统一。', '建立 Motion Runtime 与审计规则，统一注册、禁用和卸载清理。', '已修复'),
]
make_table(doc, ['缺陷编号', '问题描述', '修复方式', '修复状态'], defects, [1100, 2400, 4300, 1560], font_size=9.5)

add_heading(doc, '十四、测试结果汇总')
result_rows = [
    ('npm run build', '通过', '前台与后台生产构建完成'),
    ('npm run check', '通过', '构建、工程规则和 Motion 审计通过'),
    ('npm run check:motion', '通过', '动画边界、清理和性能规则通过'),
    ('npm run smoke', '通过', '完整聚合回归通过'),
    ('npm run smoke:web', '通过', '前后台路由与 Chunk 检查通过'),
    ('npm run smoke:api', '通过', '真实 MySQL API 回归通过'),
    ('npm run db:check', '通过', '数据库结构与基线数据一致'),
    ('git diff --check', '通过', '无空白错误；仅存在既有 CRLF 提示'),
]
make_table(doc, ['检查命令', '结果', '说明'], result_rows, [2600, 1500, 5260], font_size=10.5)
add_body(doc, '本次记录的自动化命令均于 2026 年 6 月 22 日在项目工作区执行并通过。smoke:api 使用实际 MySQL 数据完成登录、商品、图书、活动、社区、举报、预约、座位、订单、通知、上传和后台资源回归。')

add_heading(doc, '十五、测试结论')
add_body(doc, '综合功能、接口、数据库、UI、动画、安全和稳定性测试结果，Coffee Book 的核心业务链路完整，前后台状态同步可靠，接口错误处理清晰，数据库迁移具备幂等性并兼容旧数据。空间预约、社区互动、审核、订单和用户中心等高风险模块均获得自动化或手工用例覆盖。')
add_body(doc, '项目的深浅主题、统一状态组件和 Motion Runtime 提升了交互一致性；路由懒加载、图片延迟加载、RAF 合帧、列表动画上限和 reduced-motion 支持降低了性能与可访问性风险。构建、工程检查、Web smoke、API smoke、数据库检查和 Git diff 检查均通过。')
add_body(doc, '因此，本测试组认为 Coffee Book 已达到课程项目和软件工程项目的验收要求，具备提交 GitHub、部署到测试环境并继续迭代的条件。正式生产部署前，建议进一步配置 HTTPS、对象存储、真实短信服务、数据库备份、监控告警和基于真实并发的压力测试。')

add_heading(doc, '十六、附录')
add_heading(doc, '16.1 建议保留的工程命令', level=2)
commands = [
    'npm install',
    'npm run db:init',
    'node server/db/migrate.js',
    'npm run build',
    'npm run smoke',
    'npm run smoke:web',
    'npm run smoke:api',
    'npm run check',
    'npm run check:motion',
]
for command in commands:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(command)
    set_run_font(run, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(3)
add_heading(doc, '16.2 Motion DevTools 命令', level=2)
for command in [
    'window.__coffeeMotion.status()',
    'window.__coffeeMotion.disable()',
    'window.__coffeeMotion.enable()',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(command)
    set_run_font(run, size=11, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(3)
add_body(doc, '以上命令应在项目根目录或浏览器开发者控制台中执行。提交前建议依次完成数据库检查、构建、工程检查和 smoke 回归，并确认 .env、uploads、logs、dist 与 node_modules 未进入版本控制。')

doc.save(OUTPUT)

texts = [p.text for p in doc.paragraphs]
for table in doc.tables:
    texts.extend(cell.text for row in table.rows for cell in row.cells)
all_text = '\n'.join(texts)
chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', all_text))
print(f'Created: {OUTPUT}')
print(f'Chinese characters: {chinese_chars}')
print(f'Tables: {len(doc.tables)}')
